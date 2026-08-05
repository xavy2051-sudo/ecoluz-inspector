import streamlit as st
import pandas as pd
import sqlite3
import math
import io
import re
import urllib.parse
from PIL import Image as PILImage

# Librerías para exportación de documentos
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, KeepTogether, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

DB_NAME = "ecoluz_inspector.db"

# --- 1. BASE DE DATOS SQLITE ---
def init_db():
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS materiales (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT UNIQUE NOT NULL,
            unidad TEXT NOT NULL,
            rendimiento REAL NOT NULL,
            merma REAL NOT NULL,
            formato REAL NOT NULL,
            precio INTEGER NOT NULL
        )
    """)
    cursor.execute("SELECT COUNT(*) FROM materiales")
    if cursor.fetchone()[0] == 0:
        datos_iniciales = [
            ("Cerámico Muro 30x60", "m2", 1.0, 0.08, 1.44, 12500),
            ("Porcelanato Piso 60x60", "m2", 1.0, 0.10, 1.44, 18900),
            ("Adhesivo Cerámico (25kg)", "saco", 5.0, 0.05, 25.0, 8900),
            ("Frague Flexible (1kg)", "kg", 0.25, 0.05, 1.0, 2200),
            ("Pintura Esmalte al Agua", "galón", 0.08, 0.05, 1.0, 24900),
            ("Estructura Metalcom C 90x0.85", "tira", 0.35, 0.05, 6.0, 11500),
            ("Placa Yeso Cartón Volcanita 12.5mm", "m2", 1.0, 0.07, 2.88, 7800)
        ]
        cursor.executemany("""
            INSERT INTO materiales (nombre, unidad, rendimiento, merma, formato, precio)
            VALUES (?, ?, ?, ?, ?, ?)
        """, datos_iniciales)
    conn.commit()
    conn.close()

def obtener_materiales_df():
    conn = sqlite3.connect(DB_NAME)
    df = pd.read_sql_query("SELECT * FROM materiales", conn)
    conn.close()
    return df

def guardar_nuevo_material(nombre, unidad, rendimiento, merma, formato, precio):
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("""
        INSERT OR REPLACE INTO materiales (nombre, unidad, rendimiento, merma, formato, precio)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (nombre, unidad, rendimiento, merma, formato, precio))
    conn.commit()
    conn.close()

init_db()

# --- 2. FUNCIONES DE CÁLCULO Y AUXILIARES ---
def calcular_totales_eco(costo_mat, pct_mo, costo_mo_fijo, pct_gg, pct_uti, incluye_iva):
    costo_mo = (costo_mat * (pct_mo / 100.0)) + costo_mo_fijo
    costo_directo = costo_mat + costo_mo
    gastos_generales = costo_directo * (pct_gg / 100.0)
    utilidad = costo_directo * (pct_uti / 100.0)
    neto = costo_directo + gastos_generales + utilidad
    iva = neto * 0.19 if incluye_iva else 0
    total_bruto = neto + iva
    return {
        "costo_mat": costo_mat,
        "costo_mo": costo_mo,
        "costo_directo": costo_directo,
        "gg": gastos_generales,
        "utilidad": utilidad,
        "neto": neto,
        "iva": iva,
        "total_bruto": total_bruto
    }

def fmt_clp(monto):
    return f"${monto:,.0f}".replace(",", ".")

def limpiar_telefono(tel):
    return re.sub(r'[^\d]', '', str(tel))

# --- 3. GENERADORES DE INFORMES (EXCEL, WORD, PDF) ---
def generar_excel(df_cub, fotos, planos, specs, obra, inspector, fecha, eco, cond):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_meta = pd.DataFrame([
            ["PROPUESTA COMERCIAL Y DOSSIER TÉCNICO - ECOLUZ", ""],
            ["Proyecto / Obra:", obra],
            ["Inspector / Conductor:", inspector],
            ["Fecha de Inspección:", str(fecha)],
            ["", ""],
            ["RESUMEN FINANCIERO", ""],
            ["Costo Materiales:", fmt_clp(eco["costo_mat"])],
            ["Mano de Obra:", fmt_clp(eco["costo_mo"])],
            ["Costo Directo:", fmt_clp(eco["costo_directo"])],
            ["Gastos Generales:", fmt_clp(eco["gg"])],
            ["Utilidad:", fmt_clp(eco["utilidad"])],
            ["SUBTOTAL NETO:", fmt_clp(eco["neto"])],
            ["IVA (19%):", fmt_clp(eco["iva"])],
            ["TOTAL BRUTO:", fmt_clp(eco["total_bruto"])],
            ["", ""],
            ["CONDICIONES COMERCIALES", ""],
            ["Plazo de Ejecución:", f"{cond['plazo']} días hábiles"],
            ["Validez Oferta:", f"{cond['validez']} días corridos"],
            ["Forma de Pago:", cond['pago']]
        ])
        df_meta.to_excel(writer, sheet_name='Resumen Comercial', index=False, header=False)
        
        if not df_cub.empty:
            df_cub.to_excel(writer, sheet_name='Cubicaciones y APU', index=False)
            
        if fotos:
            datos_fotos = [[f["recinto"], f["titulo"], f["descripcion"], f["nombre_archivo"]] for f in fotos]
            df_fot = pd.DataFrame(datos_fotos, columns=["Recinto", "Título", "Observación Técnica", "Archivo"])
            df_fot.to_excel(writer, sheet_name='Registro Fotográfico', index=False)
            
    output.seek(0)
    return output

def generar_word(df_cub, fotos, planos, specs, obra, inspector, fecha, eco, cond):
    doc = Document()
    
    title = doc.add_heading('ECOLUZ - OFERTA ECONÓMICA Y INFORME TÉCNICO', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_info = doc.add_paragraph()
    p_info.add_run("Proyecto / Cliente: ").bold = True
    p_info.add_run(f"{obra}\n")
    p_info.add_run("Inspector Técnico: ").bold = True
    p_info.add_run(f"{inspector}\n")
    p_info.add_run("Fecha de Oferta: ").bold = True
    p_info.add_run(f"{fecha}\n")
    
    doc.add_heading('1. Resumen Financiero y Cotización', level=2)
    t_eco = doc.add_table(rows=8, cols=2)
    t_eco.style = 'Table Grid'
    filas_datos = [
        ("Materiales e Insumos", fmt_clp(eco["costo_mat"])),
        ("Mano de Obra y Equipos", fmt_clp(eco["costo_mo"])),
        ("COSTO DIRECTO TOTAL", fmt_clp(eco["costo_directo"])),
        ("Gastos Generales", fmt_clp(eco["gg"])),
        ("Utilidad", fmt_clp(eco["utilidad"])),
        ("NETO", fmt_clp(eco["neto"])),
        ("IVA (19%)", fmt_clp(eco["iva"])),
        ("TOTAL PRESUPUESTO BRUTO", fmt_clp(eco["total_bruto"]))
    ]
    for idx, (lbl, val) in enumerate(filas_datos):
        t_eco.rows[idx].cells[0].text = lbl
        t_eco.rows[idx].cells[1].text = val
        if "TOTAL" in lbl or "NETO" in lbl:
            t_eco.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            t_eco.rows[idx].cells[1].paragraphs[0].runs[0].font.bold = True
            
    doc.add_heading('2. Condiciones Comerciales', level=2)
    doc.add_paragraph(f"• Plazo de Ejecución: {cond['plazo']} días hábiles a partir de la firma/anticipo.")
    doc.add_paragraph(f"• Validez de la Cotización: {cond['validez']} días corridos.")
    doc.add_paragraph(f"• Forma de Pago: {cond['pago']}")

    doc.add_heading('3. Especificaciones Técnicas (E.T.)', level=2)
    doc.add_paragraph(specs if specs else "Sin especificaciones agregadas.")
    
    if not df_cub.empty:
        doc.add_heading('4. Detalle de Cubicaciones y APU', level=2)
        table = doc.add_table(rows=1, cols=len(df_cub.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df_cub.columns):
            hdr_cells[i].text = col_name
        for _, row in df_cub.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

    if fotos:
        doc.add_page_break()
        doc.add_heading('5. Catálogo Fotográfico de Inspección', level=2)
        for i, foto in enumerate(fotos, 1):
            doc.add_heading(f"Foto N°{i}: {foto['titulo']} ({foto['recinto']})", level=3)
            doc.add_picture(io.BytesIO(foto['bytes']), width=Inches(4.5))
            p_desc = doc.add_paragraph()
            p_desc.add_run("Observación Técnico: ").bold = True
            p_desc.add_run(foto['descripcion'])

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def preparar_imagen_pdf(img_bytes, max_width=350, max_height=250):
    img = PILImage.open(io.BytesIO(img_bytes))
    w, h = img.size
    aspect = h / float(w)
    width = max_width
    height = max_width * aspect
    if height > max_height:
        height = max_height
        width = max_height / aspect
    buf = io.BytesIO()
    img.convert('RGB').save(buf, format='JPEG', quality=85)
    buf.seek(0)
    return RLImage(buf, width=width, height=height)

def generar_pdf(df_cub, fotos, planos, specs, obra, inspector, fecha, eco, cond):
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, rightMargin=35, leftMargin=35, topMargin=35, bottomMargin=35)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=16, leading=20, textColor=colors.HexColor('#1E3A8A'))
    h2_style = ParagraphStyle('H2Style', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=11, leading=15, textColor=colors.HexColor('#1E3A8A'), spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle('BodyStyle', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=12)
    table_cell_style = ParagraphStyle('TableCellStyle', parent=styles['Normal'], fontName='Helvetica', fontSize=7.5, leading=9)
    table_hdr_style = ParagraphStyle('TableHdrStyle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.whitesmoke)

    story.append(Paragraph("ECOLUZ - PROPUESTA COMERCIAL E INFORME TÉCNICO", title_style))
    story.append(Spacer(1, 6))
    
    meta_text = f"<b>Proyecto:</b> {obra} | <b>Inspector Técnico:</b> {inspector} | <b>Fecha:</b> {fecha}"
    story.append(Paragraph(meta_text, body_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("1. Resumen Presupuestario y Oferta Económica", h2_style))
    tabla_eco_data = [
        ["Ítem / Concepto", "Valor Total"],
        ["Materiales e Insumos Directos", fmt_clp(eco["costo_mat"])],
        ["Mano de Obra y Ejecución", fmt_clp(eco["costo_mo"])],
        ["COSTO DIRECTO DE OBRA", fmt_clp(eco["costo_directo"])],
        ["Gastos Generales y Logística", fmt_clp(eco["gg"])],
        ["Utilidad de Empresa", fmt_clp(eco["utilidad"])],
        ["SUBTOTAL NETO", fmt_clp(eco["neto"])],
        ["IVA (19%)", fmt_clp(eco["iva"])],
        ["TOTAL A COTIZAR (BRUTO)", fmt_clp(eco["total_bruto"])]
    ]
    t_eco = Table(tabla_eco_data, colWidths=[280, 180])
    t_eco.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E3A8A')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
        ('FONTNAME', (0,3), (-1,3), 'Helvetica-Bold'),
        ('FONTNAME', (0,6), (-1,-1), 'Helvetica-Bold'),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#E2E8F0')),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 4),
    ]))
    story.append(t_eco)
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("2. Condiciones Comerciales", h2_style))
    cond_text = f"• <b>Plazo de ejecución:</b> {cond['plazo']} días hábiles.<br/>" \
                f"• <b>Validez de la propuesta:</b> {cond['validez']} días corridos.<br/>" \
                f"• <b>Forma de pago:</b> {cond['pago']}"
    story.append(Paragraph(cond_text, body_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph("3. Especificaciones Técnicas y Alcance", h2_style))
    texto_specs = specs.replace('\n', '<br/>') if specs else "Sin especificaciones agregadas."
    story.append(Paragraph(texto_specs, body_style))
    story.append(Spacer(1, 10))
    
    if not df_cub.empty:
        story.append(Paragraph("4. Detalle de Cubicaciones", h2_style))
        data_cub = [[Paragraph(str(col), table_hdr_style) for col in df_cub.columns]]
        for _, row in df_cub.iterrows():
            data_cub.append([Paragraph(str(v), table_cell_style) for v in row.values])
        
        t_cub = Table(data_cub, colWidths=[70, 120, 60, 60, 50, 60, 70])
        t_cub.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2563EB')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('TOPPADDING', (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ]))
        story.append(t_cub)
        
    if fotos:
        story.append(PageBreak())
        story.append(Paragraph("5. Catálogo Fotográfico", h2_style))
        for i, foto in enumerate(fotos, 1):
            elem = []
            elem.append(Paragraph(f"<b>Foto N°{i}: {foto['titulo']}</b> ({foto['recinto']})", body_style))
            elem.append(Spacer(1, 3))
            elem.append(preparar_imagen_pdf(foto['bytes']))
            elem.append(Spacer(1, 3))
            elem.append(Paragraph(f"<i>Observación:</i> {foto['descripcion']}", body_style))
            elem.append(Spacer(1, 10))
            story.append(KeepTogether(elem))

    doc.build(story)
    output.seek(0)
    return output

# --- 4. CONFIGURACIÓN PÁGINA STREAMLIT ---
st.set_page_config(page_title="ECOLUZ - Cerebro Inspector Técnico", page_icon="🏗️", layout="wide")

st.markdown("""
<style>
    .stApp { background-color: #FAFAFA; }
    .header-box { background-color: #1E3A8A; color: white; padding: 18px 24px; border-radius: 8px; margin-bottom: 20px; }
    .header-box h1 { color: white; margin: 0; font-size: 24px; }
    .header-box p { margin: 5px 0 0 0; font-size: 14px; opacity: 0.9; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="header-box">
    <h1>🏗️ ECOLUZ - Cerebro Inspector Técnico</h1>
    <p>Levantamiento, Cubicaciones, Registro Fotográfico y Cotizador Comercial Exprés</p>
</div>
""", unsafe_allow_html=True)

# --- 5. ESTADOS DE SESIÓN ---
if "cubicaciones" not in st.session_state:
    st.session_state.cubicaciones = []
if "fotos" not in st.session_state:
    st.session_state.fotos = []
if "planos" not in st.session_state:
    st.session_state.planos = []
if "especificaciones" not in st.session_state:
    st.session_state.especificaciones = """1. GENERALIDADES:
La presente especificación técnica rige los trabajos de remodelación e inspección según normativa chilena vigente NCh.

2. PREPARACIÓN DE SUSTRATOS:
Las superficies de muros y radieres deberán encontrarse limpias, secas y libres de desmoldantes antes de la aplicación de morteros o adhesivos cerámicos.

3. REVESTIMIENTOS Y TERMINACIONES:
La colocación de cerámica/porcelanato se ejecutará con doble encolado en formatos mayores a 30x30 cm. Se utilizará fragüe flexible e impermeable en zonas húmedas (baños/cocinas)."""

# --- 6. MENÚ LATERAL ---
st.sidebar.header("📋 Datos del Proyecto")
nombre_obra = st.sidebar.text_input("Nombre de la Obra / Cliente", value="Remodelación Residencial")
inspector = st.sidebar.text_input("Inspector Técnico", value="Constructor Civil")
fecha_inspeccion = st.sidebar.date_input("Fecha de Inspección")

st.sidebar.markdown("---")
st.sidebar.header("💰 Parámetros Económicos")
pct_mano_obra = st.sidebar.number_input("Mano de Obra (% sobre Materiales)", min_value=0.0, value=70.0, step=5.0)
costo_mo_fijo = st.sidebar.number_input("Mano de Obra Adicional Fija ($)", min_value=0, value=0, step=10000)
pct_gastos_generales = st.sidebar.number_input("Gastos Generales GG (%)", min_value=0.0, value=10.0, step=1.0)
pct_utilidad = st.sidebar.number_input("Utilidad (% U)", min_value=0.0, value=15.0, step=1.0)
incluye_iva = st.sidebar.checkbox("Calcular IVA (19%)", value=True)

st.sidebar.markdown("---")
st.sidebar.header("🤝 Condiciones Comerciales")
plazo_dias = st.sidebar.number_input("Plazo de Ejecución (Días)", min_value=1, value=10)
validez_dias = st.sidebar.number_input("Validez Cotización (Días)", min_value=1, value=15)
forma_pago = st.sidebar.text_input("Forma de Pago", value="50% Anticipo - 50% Recepción Conforme")

st.sidebar.markdown("---")
st.sidebar.header("📱 Contacto Cliente")
cliente_tel = st.sidebar.text_input("Teléfono WhatsApp Cliente", value="+56912345678")
cliente_email = st.sidebar.text_input("Email Cliente", value="cliente@ejemplo.com")

modulo = st.sidebar.radio(
    "Módulo de Trabajo",
    [
        "1. Levantamiento y Cubicaciones",
        "2. Registro Fotográfico y Planos",
        "3. Especificaciones Técnicas (E.T.)",
        "4. Análisis de Precios Unitarios (APU)",
        "5. Mantenedor Base de Datos (SQLite)",
        "6. Cierre Económico y Presupuesto",
        "7. Enviar Cotización a Cliente 🚀"
    ]
)

# Cálculos económicos en tiempo real
df_cubicaciones_actual = pd.DataFrame(st.session_state.cubicaciones) if len(st.session_state.cubicaciones) > 0 else pd.DataFrame()
costo_materiales_tot = df_cubicaciones_actual["Costo Total ($)"].sum() if not df_cubicaciones_actual.empty else 0
resumen_eco = calcular_totales_eco(costo_materiales_tot, pct_mano_obra, costo_mo_fijo, pct_gastos_generales, pct_utilidad, incluye_iva)
condiciones_dict = {"plazo": plazo_dias, "validez": validez_dias, "pago": forma_pago}

# --- MÓDULO 1: CUBICACIONES ---
if modulo == "1. Levantamiento y Cubicaciones":
    st.subheader("📐 Geometría y Levantamiento del Recinto")
    col1, col2 = st.columns(2)
    with col1:
        recinto = st.selectbox("Seleccione Recinto", ["Baño Principal", "Cocina", "Dormitorio", "Estar / Comedor", "Exterior / Fachada"])
        largo = st.number_input("Largo (m)", min_value=0.1, value=2.5, step=0.1)
        ancho = st.number_input("Ancho (m)", min_value=0.1, value=1.8, step=0.1)
        alto = st.number_input("Alto (m)", min_value=0.1, value=2.4, step=0.1)
    
    with col2:
        st.write("**Descuento de Vanos (Puertas / Ventanas)**")
        cant_puertas = st.number_input("Cantidad de Puertas", min_value=0, value=1)
        area_puerta = st.number_input("Superficie Promedio por Puerta (m²)", min_value=0.0, value=1.4)
        cant_ventanas = st.number_input("Cantidad de Ventanas", min_value=0, value=1)
        area_ventana = st.number_input("Superficie Promedio por Ventana (m²)", min_value=0.0, value=0.36)
        vano_total = (cant_puertas * area_puerta) + (cant_ventanas * area_ventana)

    area_piso = largo * ancho
    perimetro = 2 * (largo + ancho)
    area_muros_bruta = perimetro * alto
    area_muros_neta = max(0.0, area_muros_bruta - vano_total)

    st.markdown("---")
    st.subheader("📊 Resumen Métrico del Recinto")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Área Piso / Cielo", f"{area_piso:.2f} m²")
    m2.metric("Perímetro Muros", f"{perimetro:.2f} m")
    m3.metric("Muros Neta", f"{area_muros_neta:.2f} m²", f"-{vano_total:.2f} m² vanos")
    m4.metric("Volumen Recinto", f"{(area_piso * alto):.2f} m³")

    st.markdown("---")
    st.subheader("🛠️ Asignación de Partidas (Desde SQLite)")
    df_db = obtener_materiales_df()
    lista_materiales = df_db["nombre"].tolist()
    
    partida_seleccionada = st.selectbox("Partida / Revestimiento", lista_materiales)
    aplica_sobre = st.radio("Aplicar superficie sobre", ["Muros Netos", "Piso / Cielo"])
    superficie_aplicada = area_muros_neta if aplica_sobre == "Muros Netos" else area_piso
    
    if st.button("➕ Agregar Partida al Resumen"):
        info_mat = df_db[df_db["nombre"] == partida_seleccionada].iloc[0]
        cant_teorica = superficie_aplicada * info_mat["rendimiento"]
        cant_con_merma = cant_teorica * (1 + info_mat["merma"])
        formato_val = float(info_mat["formato"])
        
        # Corrección de cálculo de formato comercial
        unidades_comerciales = math.ceil(cant_con_merma / formato_val)
        costo_total = unidades_comerciales * int(info_mat["precio"])
        
        st.session_state.cubicaciones.append({
            "Recinto": recinto,
            "Partida": partida_seleccionada,
            "Superficie (m²)": round(superficie_aplicada, 2),
            "Cant. Merma": round(cant_con_merma, 2),
            "Unidades": unidades_comerciales,
            "Precio Unit. ($)": int(info_mat["precio"]),
            "Costo Total ($)": costo_total
        })
        st.success(f"Partida '{partida_seleccionada}' agregada.")
        st.rerun()

# --- MÓDULO 2: FOTOS Y PLANOS ---
elif modulo == "2. Registro Fotográfico y Planos":
    st.subheader("📸 Captura y Evidencia Fotográfica")
    col_up, col_list = st.columns([1, 1])
    with col_up:
        recinto_foto = st.selectbox("Recinto", ["Baño Principal", "Cocina", "Dormitorio", "Estar / Comedor", "Exterior / Fachada", "Tablero Eléctrico / Redes"])
        titulo_foto = st.text_input("Título / Elemento", placeholder="Ej: Condición de radier")
        desc_foto = st.text_area("Observación Técnica", placeholder="Ej: Presenta desnivel del 2%. Requiere mortero autonivelante.")
        archivo_foto = st.file_uploader("Fotografía", type=["jpg", "jpeg", "png"])
        
        if st.button("💾 Guardar Fotografía"):
            if archivo_foto and titulo_foto:
                st.session_state.fotos.append({
                    "recinto": recinto_foto,
                    "titulo": titulo_foto,
                    "descripcion": desc_foto,
                    "nombre_archivo": archivo_foto.name,
                    "bytes": archivo_foto.read()
                })
                st.success("Foto guardada.")
                st.rerun()
            else:
                st.warning("Debe ingresar un título y subir una imagen.")

    with col_list:
        if st.session_state.fotos:
            for idx, f in enumerate(st.session_state.fotos):
                with st.expander(f"📷 {f['recinto']} - {f['titulo']}"):
                    st.image(f['bytes'], width=250)
                    st.write(f"**Obs:** {f['descripcion']}")
                    if st.button(f"🗑️ Quitar Foto #{idx+1}", key=f"del_foto_{idx}"):
                        st.session_state.fotos.pop(idx)
                        st.rerun()

# --- MÓDULO 3: ESPECIFICACIONES ---
elif modulo == "3. Especificaciones Técnicas (E.T.)":
    st.subheader("📝 Redacción de Especificaciones Técnicas (E.T.)")
    st.session_state.especificaciones = st.text_area("Especificaciones del Proyecto", value=st.session_state.especificaciones, height=300)

# --- MÓDULO 4: APU ---
elif modulo == "4. Análisis de Precios Unitarios (APU)":
    st.subheader("📦 Detalle de Materiales Cubicados")
    if not df_cubicaciones_actual.empty:
        st.dataframe(df_cubicaciones_actual, use_container_width=True)
        st.metric("Total Materiales", fmt_clp(costo_materiales_tot))
        
        col_del1, col_del2 = st.columns([1, 4])
        with col_del1:
            if st.button("🗑️ Limpiar Toda la Lista"):
                st.session_state.cubicaciones = []
                st.rerun()
    else:
        st.info("No hay partidas ingresadas.")

# --- MÓDULO 5: SQLITE ---
elif modulo == "5. Mantenedor Base de Datos (SQLite)":
    st.subheader("🗄️ Mantenedor de Insumos")
    col_add, col_view = st.columns([1, 2])
    with col_add:
        with st.form("form_mat"):
            nom = st.text_input("Nombre Material / Insumo")
            uni = st.selectbox("Unidad", ["m2", "saco", "kg", "galón", "tira", "unidad"])
            rend = st.number_input("Rendimiento", min_value=0.01, value=1.0)
            mer = st.number_input("Merma", min_value=0.0, value=0.05)
            fmt = st.number_input("Formato", min_value=0.01, value=1.0)
            prc = st.number_input("Precio ($)", min_value=1, value=5000)
            if st.form_submit_button("💾 Guardar"):
                if nom.strip():
                    guardar_nuevo_material(nom.strip(), uni, rend, mer, fmt, prc)
                    st.success("Material guardado correctamente.")
                    st.rerun()
                else:
                    st.error("Ingrese un nombre válido.")
    with col_view:
        st.dataframe(obtener_materiales_df(), use_container_width=True)

# --- MÓDULO 6: CIERRE ECONÓMICO ---
elif modulo == "6. Cierre Económico y Presupuesto":
    st.subheader("📊 Estructura de Costos de la Propuesta")
    
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Materiales", fmt_clp(resumen_eco["costo_mat"]))
    c2.metric("Mano de Obra", fmt_clp(resumen_eco["costo_mo"]), f"{pct_mano_obra}% mat.")
    c3.metric("Costo Directo", fmt_clp(resumen_eco["costo_directo"]))
    c4.metric("Gastos Gen. + Utilidad", fmt_clp(resumen_eco["gg"] + resumen_eco["utilidad"]))
    
    st.markdown("---")
    c_neto, c_iva, c_total = st.columns(3)
    c_neto.metric("Subtotal Neto", fmt_clp(resumen_eco["neto"]))
    c_iva.metric("IVA (19%)", fmt_clp(resumen_eco["iva"]))
    c_total.metric("TOTAL A COTIZAR (BRUTO)", fmt_clp(resumen_eco["total_bruto"]))

# --- MÓDULO 7: ENVIAR COTIZACIÓN A CLIENTE 🚀 ---
elif modulo == "7. Enviar Cotización a Cliente 🚀":
    st.subheader("🚀 Enviar Cotización Formal al Cliente")
    st.markdown(f"**Cliente / Proyecto:** `{nombre_obra}` | **Inspector:** `{inspector}`")
    
    st.info(f"💰 **VALOR TOTAL BRUTO DE LA PROPUESTA:** {fmt_clp(resumen_eco['total_bruto'])}")
    
    st.markdown("---")
    st.markdown("### 📤 Canales de Envío Rápido")
    
    msg_whatsapp = (
        f"Hola! Le envío la cotización para el proyecto *{nombre_obra}*:\n\n"
        f"🔹 *Costo Materiales:* {fmt_clp(resumen_eco['costo_mat'])}\n"
        f"🔹 *Mano de Obra:* {fmt_clp(resumen_eco['costo_mo'])}\n"
        f"🔹 *Subtotal Neto:* {fmt_clp(resumen_eco['neto'])}\n"
        f"🔹 *IVA (19%):* {fmt_clp(resumen_eco['iva'])}\n"
        f"✅ *TOTAL PRESUPUESTO:* {fmt_clp(resumen_eco['total_bruto'])}\n\n"
        f"📌 *Condiciones:*\n"
        f"- Plazo de Ejecución: {plazo_dias} días hábiles\n"
        f"- Validez: {validez_dias} días corridos\n"
        f"- Pago: {forma_pago}\n\n"
        f"Le adjunto el informe técnico completo en PDF/Word. Saludos!"
    )
    
    tel_clean = limpiar_telefono(cliente_tel)
    url_whatsapp = f"https://wa.me/{tel_clean}?text={urllib.parse.quote(msg_whatsapp)}"
    
    subject_email = f"Cotización Proyecto - {nombre_obra}"
    body_email = msg_whatsapp.replace('\n', '%0D%0A')
    url_email = f"mailto:{cliente_email}?subject={urllib.parse.quote(subject_email)}&body={body_email}"

    col_wa, col_mail = st.columns(2)
    
    with col_wa:
        st.markdown("#### 📲 Envío por WhatsApp")
        st.write(f"Destinatario: `{cliente_tel}`")
        st.link_button("📲 Abrir WhatsApp y Enviar Resumen", url_whatsapp, type="primary")
        st.caption("Abre la app de WhatsApp en el celular o WhatsApp Web en PC con el mensaje precargado.")

    with col_mail:
        st.markdown("#### ✉️ Envío por Correo Electrónico")
        st.write(f"Destinatario: `{cliente_email}`")
        st.link_button("✉️ Abrir Cliente de Correo", url_email)
        st.caption("Abre tu aplicación de correo predeterminada (Outlook, Gmail, Apple Mail).")

    st.markdown("---")
    st.markdown("### 📥 Descargar Documentos para Adjuntar al Cliente")
    st.write("Descarga el archivo a tu celular o PC y adjúntalo al mensaje de WhatsApp o Correo:")
    
    col_d1, col_d2, col_d3 = st.columns(3)
    nombre_archivo_base = re.sub(r'\W+', '_', nombre_obra).lower()
    
    with col_d1:
        pdf_bytes = generar_pdf(df_cubicaciones_actual, st.session_state.fotos, st.session_state.planos, st.session_state.especificaciones, nombre_obra, inspector, fecha_inspeccion, resumen_eco, condiciones_dict)
        st.download_button(
            label="📑 Descargar PDF Oficial (.pdf)",
            data=pdf_bytes,
            file_name=f"cotizacion_{nombre_archivo_base}.pdf",
            mime="application/pdf",
            type="primary"
        )
        
    with col_d2:
        word_bytes = generar_word(df_cubicaciones_actual, st.session_state.fotos, st.session_state.planos, st.session_state.especificaciones, nombre_obra, inspector, fecha_inspeccion, resumen_eco, condiciones_dict)
        st.download_button(
            label="📝 Descargar Word (.docx)",
            data=word_bytes,
            file_name=f"cotizacion_{nombre_archivo_base}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    with col_d3:
        excel_bytes = generar_excel(df_cubicaciones_actual, st.session_state.fotos, st.session_state.planos, st.session_state.especificaciones, nombre_obra, inspector, fecha_inspeccion, resumen_eco, condiciones_dict)
        st.download_button(
            label="📊 Descargar Excel (.xlsx)",
            data=excel_bytes,
            file_name=f"cotizacion_{nombre_archivo_base}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )